import docx.shared
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from setup import course_list
import os

class CourseBook():
    def __init__(self, excel_path):
        self.excel_path = excel_path
        self.create_dir()
        self.result = self.parser_excel()
        for course in self.result.items():
            self.analyze_data(course)
        self.write_file()

    def create_dir(self):
        path = os.getcwd()
        self.word_path = os.path.join(path, "选课小本本.docx")
        self.dir_path = os.path.join(path, "pic")
        try:
            os.makedirs(self.dir_path, exist_ok=True)
            print(f"目录 '{self.dir_path}' 已创建或已存在。")
        except OSError as error:
            print(f"创建目录 '{self.dir_path}' 失败：{error}")

    def parser_excel(self):
        # 1、读取excel文件；
        # 2、按照课程分类：{course: [{'teacher':'', 'score':[], 'note':'', 'time':''}]}]
        # 3、替换课程（excel中课程以数字代替，因此替换）
        data = pd.read_excel(self.excel_path)
        result = {}
        for item in data.iterrows():
            course = course_list[item[1][6]][1]
            teacher = item[1][7]
            score = [item[1][i] for i in range(8,13)]
            note = item[1][13]
            time = item[1][1]
            if course not in result.keys():
                result[course] = []
            result[course].append({'teacher':teacher, 'score':score, 'note':note, 'time':time})
        print("数据处理完毕")
        return result

    def analyze_data(self, course):
        def SumData(data):
            # 初始化计数数组，长度为评价等级的数量，初始值为0
            count_array = [[0,0,0,0,0],[0,0,0,0,0],[0,0,0,0,0],[0,0,0,0,0],[0,0,0,0,0]]
            for i, scores in enumerate(data):
                for j, score in enumerate(scores):
                    if i < 4:
                        count_array[i][score - 1] += 1
                    else:
                        count_array[i][int(score/2) - 1] += 1
            return count_array

        def Picture_Draw(data, data_cum, name):
            labels = ['知识含金量', '给分情况', '难易程度', '友善程度', '推荐意愿']
            degree = ["Strongly Bad", "Bad", "Neither", "Good", "Very good"]
            category_colors = plt.cm.RdYlGn(np.linspace(0.05, 0.95, len(degree)))
            plt.rcParams['font.sans-serif'] = ['SimHei']
            plt.rcParams.update({'font.size': 16})

            fig, ax = plt.subplots(figsize=(19.2, 10.8))
            ax.invert_yaxis()
            ax.xaxis.set_visible(False)
            ax.set_xlim(0, np.sum(data, axis=1).max())

            # TODO 这一段是加标签，但是需要修改
            for i, (colname, color) in enumerate(zip(degree, category_colors)):
                widths = data[:, i]
                starts = data_cum[:, i] - widths
                rects = ax.barh(labels, widths, left=starts, height=0.5, label=colname, color=color)

                # 计算文本颜色
                r, g, b, _ = color
                text_color = 'white' if r * g * b < 0.5 else 'darkgrey'

                # 使用 ax.text() 显示标签
                for rect in rects:
                    width = rect.get_width()  # 获取条形的宽度
                    height = rect.get_height()  # 获取条形的高度
                    ax.text(width + starts[i], rect.get_y() + height / 2,  # 计算文本位置
                            str(int(width)), ha='center', va='center', color=text_color)

            ax.legend(ncols=len(degree), bbox_to_anchor=(0, 1), loc='lower left', fontsize='small')
            plt.title(name, loc="right")
            name_file = f"{name}.jpg"
            plt.savefig(f"{os.path.join(self.dir_path, name_file)}")
            print(f"图像准备完毕 {name}")
            # plt.show()

        course_name = course[0]
        course_data = course[1]
        scores_list = np.array([item['score'] for item in course_data]).T
        count_array = np.array(SumData(scores_list))  # 统计每个分数的数量
        count_array_cum = count_array.cumsum(axis=1)  # 对数量做累积，为之后画图做准备
        Picture_Draw(count_array, count_array_cum, course_name)

    def write_file(self):

        doc = Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        doc.add_page_break()
        doc.save(self.word_path)

        for i, (name, judge_list) in enumerate(self.result.items()):
            doc.add_heading(name, level = 1)
            try:
                name_file = f"{name}.jpg"
                doc.add_picture(f"{os.path.join(self.dir_path, name_file)}", height=docx.shared.Cm(8.5))
            except:
                doc.add_paragraph("没有图片展示")
            for judge_dict in judge_list:
                paragraph = f"· {judge_dict['time']} \n" \
                            f"\t· 老师：{judge_dict['teacher']} \n" \
                            f"\t· 评价：{judge_dict['note']}"
                doc.add_paragraph(paragraph)
            doc.add_page_break()
            doc.save(self.word_path)
            print(f"已经准备完毕 {name}")

if __name__ == "__main__":
    excel_path = "小本本内容.xlsx"
    CourseBook(excel_path=excel_path)
    print(f"\n全部完成，请阅读README.pdf文件完成后续操作")